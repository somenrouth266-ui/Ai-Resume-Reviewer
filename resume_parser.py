"""
resume_parser.py — Resume File Parsing Module
===============================================
Extracts plain text from PDF and DOCX resume files.
Uses pdfplumber for text-based PDFs.
Falls back to PyMuPDF (fitz) for complex/designed PDFs.
Images and graphics are silently ignored in all cases.
"""

import io
import pdfplumber
from docx import Document
from docx.oxml.ns import qn


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF using multiple strategies.

    Strategy 1: pdfplumber  — works for most standard PDFs
    Strategy 2: PyMuPDF     — works for designed/complex PDFs with
                               text boxes, columns, and embedded fonts
    Strategy 3: Raw string  — last resort character-level extraction

    Images and photos are silently ignored in all strategies.

    Args:
        file_bytes: Raw bytes of the uploaded PDF file.

    Returns:
        Extracted text string.

    Raises:
        ValueError: If all strategies fail to extract any text.
    """

    # ── Strategy 1: pdfplumber ────────────────────────────────────────────────
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                except Exception:
                    continue
        result = "\n\n".join(text_parts).strip()
        if result and len(result) > 50:
            return result
    except Exception:
        pass

    # ── Strategy 2: PyMuPDF (fitz) ───────────────────────────────────────────
    # Handles designed resumes, multi-column layouts, text boxes, custom fonts
    try:
        import fitz  # PyMuPDF
        text_parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            try:
                # get_text("text") extracts plain text, ignoring images
                page_text = page.get_text("text")
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            except Exception:
                continue
        doc.close()
        result = "\n\n".join(text_parts).strip()
        if result and len(result) > 50:
            return result
    except ImportError:
        pass  # PyMuPDF not installed, continue to next strategy
    except Exception:
        pass

    # ── Strategy 3: PyMuPDF HTML extraction ──────────────────────────────────
    # Some PDFs store text in HTML-like blocks — extract and strip tags
    try:
        import fitz
        import re
        text_parts = []
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            try:
                html = page.get_text("html")
                # Strip HTML tags to get plain text
                clean = re.sub(r"<[^>]+>", " ", html)
                clean = re.sub(r"\s+", " ", clean).strip()
                if clean and len(clean) > 20:
                    text_parts.append(clean)
            except Exception:
                continue
        doc.close()
        result = "\n\n".join(text_parts).strip()
        if result and len(result) > 50:
            return result
    except Exception:
        pass

    # ── Strategy 4: pdfplumber loose settings ────────────────────────────────
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception:
                    continue
        result = "\n\n".join(text_parts).strip()
        if result and len(result) > 50:
            return result
    except Exception:
        pass

    # ── All strategies failed ─────────────────────────────────────────────────
    raise ValueError(
        "Could not extract text from this PDF.\n\n"
        "This usually happens when:\n"
        "• The resume was saved as an image inside a PDF\n"
        "• The PDF uses heavily encrypted or non-standard fonts\n\n"
        "Please try:\n"
        "1. Open your resume in Word or Google Docs\n"
        "2. Export/Download as PDF again\n"
        "3. Upload the new PDF"
    )


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file provided as raw bytes.
    Images, embedded objects, and drawing elements are silently skipped.
    Text from paragraphs, tables, text boxes, and headers is included.

    Args:
        file_bytes: Raw bytes of the uploaded DOCX file.

    Returns:
        A single string containing all extracted text.

    Raises:
        ValueError: If the bytes cannot be parsed as a valid DOCX document.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # ── Main body paragraphs ──────────────────────────────────────────────
        for para in doc.paragraphs:
            try:
                text = para.text.strip()
                if text:
                    parts.append(text)
            except Exception:
                continue

        # ── Tables ────────────────────────────────────────────────────────────
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    try:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    except Exception:
                        continue
                if row_texts:
                    parts.append("  ".join(row_texts))

        # ── Headers and footers ───────────────────────────────────────────────
        for section in doc.sections:
            try:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
            except Exception:
                pass
            try:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
            except Exception:
                pass

        # ── Text boxes ────────────────────────────────────────────────────────
        try:
            body = doc.element.body
            for txbx in body.iter(qn("w:txbxContent")):
                for para in txbx.iter(qn("w:p")):
                    texts = [
                        node.text
                        for node in para.iter(qn("w:t"))
                        if node.text
                    ]
                    line = "".join(texts).strip()
                    if line:
                        parts.append(line)
        except Exception:
            pass

        result = "\n".join(parts).strip()
        if not result:
            raise ValueError("No text could be extracted from this DOCX file.")
        return result

    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Failed to parse DOCX: {exc}") from exc
